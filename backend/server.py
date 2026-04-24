from fastapi import FastAPI, APIRouter, UploadFile, File, Depends, HTTPException
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
import json
import uuid
import asyncio
import jwt as pyjwt
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone, timedelta

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook, Workbook
from PIL import Image
import pytesseract

from emergentintegrations.llm.chat import LlmChat, UserMessage


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
APP_PIN = os.environ.get('APP_PIN', '1497412781')
JWT_SECRET = os.environ.get('JWT_SECRET', 'dev-secret')
JWT_ALGO = 'HS256'
JWT_EXPIRES_HOURS = 24 * 7  # 1 week
CLAUDE_MODEL = 'claude-sonnet-4-5-20250929'

app = FastAPI(title="Insurance Benefits Summarizer")
api_router = APIRouter(prefix="/api")
security = HTTPBearer(auto_error=False)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
)
logger = logging.getLogger(__name__)


# ---------------- Models ----------------
class PinRequest(BaseModel):
    pin: str


class TokenResponse(BaseModel):
    token: str
    expires_at: str


class BenefitLineData(BaseModel):
    model_config = ConfigDict(extra="allow")
    plan_design: List[Dict[str, Any]] = Field(default_factory=list)
    rates: List[Dict[str, Any]] = Field(default_factory=list)
    coverages: List[Dict[str, Any]] = Field(default_factory=list)
    notes: Optional[str] = None


class Analysis(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    source_files: List[Dict[str, Any]] = Field(default_factory=list)
    executive_summary: str = ""
    carriers: List[str] = Field(default_factory=list)
    effective_date: Optional[str] = None
    benefits: Dict[str, BenefitLineData] = Field(default_factory=dict)


class AnalysisSummary(BaseModel):
    id: str
    title: str
    created_at: datetime
    carriers: List[str] = []
    file_count: int = 0


# ---------------- Auth helpers ----------------
def create_token() -> TokenResponse:
    expires = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRES_HOURS)
    payload = {"sub": "insurance-app-user", "exp": expires}
    token = pyjwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGO)
    return TokenResponse(token=token, expires_at=expires.isoformat())


def require_auth(creds: Optional[HTTPAuthorizationCredentials] = Depends(security)) -> str:
    if creds is None or not creds.credentials:
        raise HTTPException(status_code=401, detail="Missing token")
    try:
        payload = pyjwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALGO])
        return payload.get("sub", "")
    except pyjwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except pyjwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


# ---------------- File extractors ----------------
def extract_pdf(data: bytes) -> str:
    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_text_chunks = []
            for t in tables:
                for row in t:
                    if row:
                        table_text_chunks.append("\t".join([str(c) if c is not None else "" for c in row]))
            page_chunk = f"--- Page {i+1} ---\n{text}"
            if table_text_chunks:
                page_chunk += "\n[TABLES]\n" + "\n".join(table_text_chunks)
            parts.append(page_chunk)
    return "\n\n".join(parts)


def extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    out: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            out.append("\t".join(cells))
    return "\n".join(out)


def extract_xlsx(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), data_only=True)
    out: List[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() != "" for c in row):
                out.append("\t".join(["" if c is None else str(c) for c in row]))
    return "\n".join(out)


def extract_image(data: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img)
        return text.strip() or "[Image uploaded — no OCR text detected]"
    except Exception as e:
        return f"[Image processing error: {e}]"


def extract_file_text(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf(content)
    if name.endswith(".docx"):
        return extract_docx(content)
    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        return extract_xlsx(content)
    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff")):
        return extract_image(content)
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return f"[Unsupported file type: {filename}]"


# ---------------- LLM Extraction ----------------
EXTRACTION_SYSTEM_PROMPT = """You are an expert insurance benefits analyst. You will receive raw text extracted from one or more benefits-related documents (plan summaries, rate sheets, benefit summaries, SBCs, carrier proposals).

Your job: extract structured insurance benefit data for EXACTLY these 6 lines of coverage:
- dental
- vision
- basic_life
- voluntary_life
- std  (Short-Term Disability)
- ltd  (Long-Term Disability)

For each line that appears in the documents, extract three categories:

1. plan_design: array of {"label": string, "value": string} rows capturing in-network/out-of-network structure, deductibles, coinsurance, copays, frequencies, benefit maximums, waiting periods, elimination periods, benefit durations, etc.

2. rates: array of {"tier": string, "rate": string, "frequency": string (e.g. "monthly", "weekly"), "notes": string (optional)} covering tiered premium/rate structures (EE, EE+Spouse, EE+Child(ren), Family, or age-banded rates for voluntary life). Keep currency symbols.

3. coverages: array of {"label": string, "value": string} capturing benefit amounts, % of salary replaced, max benefits, guarantee issue amounts, AD&D, etc.

Also extract:
- executive_summary: 3-5 sentence plain-English overview of the ENTIRE benefit package across all lines, written for an executive audience.
- carriers: array of carrier names found (e.g. ["Guardian", "Unum"]).
- effective_date: string if found, else null.
- title: short descriptive title like "ACME Corp 2025 Benefits Renewal" (infer from content).

OUTPUT STRICT JSON with this shape (no prose, no markdown fences):

{
  "title": "...",
  "executive_summary": "...",
  "carriers": ["..."],
  "effective_date": "..." | null,
  "benefits": {
    "dental": {"plan_design": [...], "rates": [...], "coverages": [...], "notes": "..."},
    "vision": {...},
    "basic_life": {...},
    "voluntary_life": {...},
    "std": {...},
    "ltd": {...}
  }
}

If a benefit line is NOT in the documents, include it with empty arrays and notes: "Not found in provided documents." Do NOT invent data."""


async def run_llm_extraction(combined_text: str) -> Dict[str, Any]:
    if not EMERGENT_LLM_KEY:
        raise HTTPException(status_code=500, detail="EMERGENT_LLM_KEY not configured")

    session_id = f"insurance-{uuid.uuid4()}"
    chat = (
        LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=session_id,
            system_message=EXTRACTION_SYSTEM_PROMPT,
        )
        .with_model("anthropic", CLAUDE_MODEL)
        .with_params(timeout=240, num_retries=1, max_tokens=8000)
    )

    # Keep text within reasonable size. Claude context handles more,
    # but we want snappy responses (<90s) so we trim aggressively.
    MAX_CHARS = 60_000
    if len(combined_text) > MAX_CHARS:
        combined_text = combined_text[:MAX_CHARS] + "\n\n[... truncated for length ...]"

    msg = UserMessage(
        text=(
            "Extract the structured insurance benefit data from the following documents. "
            "Return ONLY valid JSON, no prose, no code fences.\n\n"
            f"=== DOCUMENTS ===\n{combined_text}\n=== END ==="
        )
    )
    try:
        response = await asyncio.wait_for(chat.send_message(msg), timeout=240)
    except asyncio.TimeoutError:
        logger.error("LLM call timed out after 240s")
        raise HTTPException(
            status_code=504,
            detail="The AI took too long to respond. Try uploading a smaller document or fewer files.",
        )
    except Exception as e:
        logger.exception(f"LLM call failed: {e}")
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {e}")

    # Clean potential markdown fences
    text = response.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
        text = text.strip()
    # Find first { and last }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1:
        text = text[start : end + 1]

    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error: {e}\nRaw: {response[:500]}")
        raise HTTPException(status_code=502, detail=f"LLM returned invalid JSON: {e}")


# ---------------- Routes ----------------
@api_router.get("/")
async def root():
    return {"message": "Insurance Benefits Summarizer API", "status": "ok"}


@api_router.post("/auth/pin", response_model=TokenResponse)
async def auth_pin(req: PinRequest):
    if req.pin != APP_PIN:
        raise HTTPException(status_code=401, detail="Invalid PIN")
    return create_token()


@api_router.get("/auth/me")
async def auth_me(user: str = Depends(require_auth)):
    return {"user": user, "authenticated": True}


BENEFIT_KEYS = ["dental", "vision", "basic_life", "voluntary_life", "std", "ltd"]


def _normalize_benefits(data: Dict[str, Any]) -> Dict[str, Any]:
    benefits = data.get("benefits") or {}
    for k in BENEFIT_KEYS:
        entry = benefits.get(k) or {}
        benefits[k] = {
            "plan_design": entry.get("plan_design") or [],
            "rates": entry.get("rates") or [],
            "coverages": entry.get("coverages") or [],
            "notes": entry.get("notes") or "",
        }
    data["benefits"] = benefits
    return data


@api_router.post("/analyze")
async def analyze(
    files: List[UploadFile] = File(...),
    user: str = Depends(require_auth),
):
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    source_files: List[Dict[str, Any]] = []
    combined_parts: List[str] = []

    for f in files:
        content = await f.read()
        size = len(content)
        try:
            text = await asyncio.to_thread(extract_file_text, f.filename or "file", content)
        except Exception as e:
            logger.exception(f"Failed to extract {f.filename}: {e}")
            raise HTTPException(
                status_code=400,
                detail=f"Could not read '{f.filename}': {e}. Try a different format.",
            )
        source_files.append({
            "name": f.filename or "file",
            "size": size,
            "content_type": f.content_type or "application/octet-stream",
            "extracted_chars": len(text),
        })
        combined_parts.append(f"=== FILE: {f.filename} ===\n{text}")

    combined_text = "\n\n".join(combined_parts)

    if not combined_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No readable text could be extracted from the uploaded files. Images may be scans without OCR-friendly text.",
        )

    logger.info(
        f"Analyzing {len(files)} file(s), {len(combined_text):,} chars extracted, user={user}"
    )
    llm_data = await run_llm_extraction(combined_text)
    llm_data = _normalize_benefits(llm_data)

    analysis_id = str(uuid.uuid4())
    created_at = datetime.now(timezone.utc)

    doc = {
        "id": analysis_id,
        "title": llm_data.get("title") or f"Analysis {created_at.strftime('%Y-%m-%d %H:%M')}",
        "created_at": created_at.isoformat(),
        "source_files": source_files,
        "executive_summary": llm_data.get("executive_summary") or "",
        "carriers": llm_data.get("carriers") or [],
        "effective_date": llm_data.get("effective_date"),
        "benefits": llm_data["benefits"],
    }

    await db.analyses.insert_one(dict(doc))
    doc.pop("_id", None)
    return doc


@api_router.get("/analyses", response_model=List[AnalysisSummary])
async def list_analyses(user: str = Depends(require_auth)):
    rows = await db.analyses.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    summaries: List[AnalysisSummary] = []
    for r in rows:
        created = r.get("created_at")
        if isinstance(created, str):
            try:
                created_dt = datetime.fromisoformat(created)
            except ValueError:
                created_dt = datetime.now(timezone.utc)
        else:
            created_dt = created or datetime.now(timezone.utc)
        summaries.append(AnalysisSummary(
            id=r["id"],
            title=r.get("title", "Untitled"),
            created_at=created_dt,
            carriers=r.get("carriers", []),
            file_count=len(r.get("source_files", [])),
        ))
    return summaries


@api_router.get("/analyses/{analysis_id}")
async def get_analysis(analysis_id: str, user: str = Depends(require_auth)):
    row = await db.analyses.find_one({"id": analysis_id}, {"_id": 0})
    if not row:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return row


@api_router.delete("/analyses/{analysis_id}")
async def delete_analysis(analysis_id: str, user: str = Depends(require_auth)):
    result = await db.analyses.delete_one({"id": analysis_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {"deleted": True}


@api_router.get("/analyses/{analysis_id}/export")
async def export_analysis(analysis_id: str, token: Optional[str] = None):
    # Allow token via query string for browser download
    if not token:
        raise HTTPException(status_code=401, detail="Missing token")
    try:
        pyjwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGO])
    except pyjwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except pyjwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

    row = await db.analyses.find_one({"id": analysis_id}, {"_id": 0})
    if not row:
        raise HTTPException(status_code=404, detail="Analysis not found")

    wb = Workbook()
    # Overview sheet
    ws = wb.active
    ws.title = "Overview"
    ws.append(["Title", row.get("title", "")])
    ws.append(["Created", str(row.get("created_at", ""))])
    ws.append(["Effective Date", row.get("effective_date") or ""])
    ws.append(["Carriers", ", ".join(row.get("carriers", []))])
    ws.append([])
    ws.append(["Executive Summary"])
    ws.append([row.get("executive_summary", "")])
    ws.append([])
    ws.append(["Source Files"])
    ws.append(["Name", "Size (bytes)", "Type"])
    for f in row.get("source_files", []):
        ws.append([f.get("name", ""), f.get("size", 0), f.get("content_type", "")])

    benefit_labels = {
        "dental": "Dental",
        "vision": "Vision",
        "basic_life": "Basic Life",
        "voluntary_life": "Voluntary Life",
        "std": "STD",
        "ltd": "LTD",
    }

    benefits = row.get("benefits", {})
    for key, label in benefit_labels.items():
        entry = benefits.get(key, {})
        sheet = wb.create_sheet(title=label[:31])
        sheet.append([f"{label} — Plan Design"])
        sheet.append(["Label", "Value"])
        for item in entry.get("plan_design", []):
            sheet.append([item.get("label", ""), item.get("value", "")])
        sheet.append([])
        sheet.append([f"{label} — Rates"])
        sheet.append(["Tier", "Rate", "Frequency", "Notes"])
        for r in entry.get("rates", []):
            sheet.append([r.get("tier", ""), r.get("rate", ""), r.get("frequency", ""), r.get("notes", "")])
        sheet.append([])
        sheet.append([f"{label} — Coverages"])
        sheet.append(["Label", "Value"])
        for item in entry.get("coverages", []):
            sheet.append([item.get("label", ""), item.get("value", "")])
        if entry.get("notes"):
            sheet.append([])
            sheet.append(["Notes"])
            sheet.append([entry.get("notes", "")])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in row.get("title", "analysis") if c.isalnum() or c in (" ", "-", "_")).strip().replace(" ", "_")[:60] or "analysis"
    filename = f"{safe_title}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
